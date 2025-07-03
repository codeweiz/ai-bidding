#!/usr/bin/env python3
"""
演示模板样式功能
"""

from pathlib import Path
from docx import Document


def demo_template_styles():
    """演示模板样式应用"""
    print("🎨 AI投标系统 - 模板样式演示")
    print("=" * 50)
    
    template_path = Path("tests/data/投标文件template.docx")
    if not template_path.exists():
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    # 加载模板
    doc = Document(str(template_path))
    
    # 清空现有内容
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
    
    print("📝 正在创建演示文档...")
    
    # 添加标题
    title = doc.add_paragraph("IPTV智慧广电系统建设项目技术方案")
    title.style = "标书1级"
    
    # 添加一级章节
    section1 = doc.add_paragraph("1. 系统总体设计")
    section1.style = "标书2级"
    
    content1 = doc.add_paragraph(
        "本项目旨在构建一套完整的IPTV智慧广电系统，采用先进的云原生架构，"
        "实现高可用、高并发、高扩展的技术目标。系统将充分满足招标文件中"
        "提出的各项技术要求和功能需求。"
    )
    content1.style = "标书正文"
    
    # 添加二级章节
    subsection1 = doc.add_paragraph("1.1 系统架构设计")
    subsection1.style = "标书3级"
    
    content2 = doc.add_paragraph(
        "系统采用微服务架构设计，将复杂的业务功能拆分为多个独立的服务单元。"
        "每个服务单元具备独立的数据存储、业务逻辑和接口定义，通过标准化的"
        "API接口进行服务间通信。"
    )
    content2.style = "标书正文"
    
    # 添加三级章节
    subsubsection1 = doc.add_paragraph("1.1.1 微服务架构设计")
    subsubsection1.style = "标书4级"
    
    content3 = doc.add_paragraph(
        "微服务架构采用Spring Cloud技术栈，包含服务注册与发现、配置管理、"
        "负载均衡、熔断降级等核心组件。各服务通过Docker容器化部署，"
        "支持弹性伸缩和故障自愈。"
    )
    content3.style = "标书正文"
    
    # 添加四级章节
    subsubsubsection1 = doc.add_paragraph("1.1.1.1 服务注册与发现")
    subsubsubsection1.style = "标书5级"
    
    content4 = doc.add_paragraph(
        "采用Eureka作为服务注册中心，实现服务的自动注册与发现。"
        "服务启动时自动向注册中心注册服务信息，客户端通过注册中心"
        "获取可用服务列表，实现动态服务调用。"
    )
    content4.style = "标书正文"
    
    # 添加第二个一级章节
    section2 = doc.add_paragraph("2. 功能设计方案")
    section2.style = "标书2级"
    
    content5 = doc.add_paragraph(
        "系统功能设计严格按照招标文件要求，涵盖内容管理、用户管理、"
        "设备管理、监控运维等核心功能模块。每个功能模块都具备完整的"
        "业务流程和用户交互界面。"
    )
    content5.style = "标书正文"
    
    # 添加技术架构图说明
    subsection2 = doc.add_paragraph("2.1 技术架构图")
    subsection2.style = "标书3级"
    
    content6 = doc.add_paragraph(
        "系统技术架构如下图所示："
    )
    content6.style = "标书正文"
    
    # 添加mermaid代码块（演示）
    mermaid_code = doc.add_paragraph("""
```mermaid
graph TD
    A[用户层] --> B[网关层]
    B --> C[业务服务层]
    C --> D[数据访问层]
    D --> E[数据存储层]
    
    subgraph "业务服务层"
        C1[内容管理服务]
        C2[用户管理服务]
        C3[设备管理服务]
        C4[监控服务]
    end
```
    """)
    mermaid_code.style = "标书正文"
    
    # 保存演示文档
    output_path = Path("outputs/demo_template_styles.docx")
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✅ 演示文档已生成: {output_path}")
    print("\n📋 使用的样式:")
    print("  - 标书1级: 主标题")
    print("  - 标书2级: 一级章节")
    print("  - 标书3级: 二级章节")
    print("  - 标书4级: 三级章节")
    print("  - 标书5级: 四级章节")
    print("  - 标书正文: 正文内容")
    
    print("\n🎯 特点:")
    print("  ✅ 自动应用预定义样式")
    print("  ✅ 保持格式一致性")
    print("  ✅ 支持多级标题层次")
    print("  ✅ 包含技术架构图代码")
    
    return output_path


def check_template_styles():
    """检查模板中的可用样式"""
    print("\n🔍 检查模板样式...")
    
    template_path = Path("tests/data/投标文件template.docx")
    if not template_path.exists():
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    doc = Document(str(template_path))
    
    print("📋 模板中的标书样式:")
    bidding_styles = []
    for style in doc.styles:
        if "标书" in style.name and "字符" not in style.name:
            bidding_styles.append(style.name)
            print(f"  ✅ {style.name}")
    
    print(f"\n📊 共找到 {len(bidding_styles)} 个标书样式")
    return bidding_styles


if __name__ == "__main__":
    # 检查模板样式
    check_template_styles()
    
    # 生成演示文档
    demo_template_styles()
    
    print("\n" + "=" * 50)
    print("🎉 演示完成!")
    print("📄 请查看生成的演示文档以了解样式效果")
