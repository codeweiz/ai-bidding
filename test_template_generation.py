#!/usr/bin/env python3
"""
测试模板样式的投标书生成
"""

import asyncio
import sys
from pathlib import Path

# 添加项目根目录到Python路径
sys.path.append(str(Path(__file__).parent))

from backend.models.project import Project
from backend.services.content_generator import ContentGenerator
from backend.models.generation import WorkflowState


async def test_template_generation():
    """测试使用模板样式生成投标书"""
    print("🧪 测试模板样式投标书生成...")
    
    # 创建测试项目
    project = Project(
        id="test_template_001",
        name="IPTV系统建设项目",
        description="测试模板样式生成",
        enable_differentiation=True
    )
    
    # 准备测试文档路径
    test_doc_path = "uploads/test_bidding_document.txt"
    template_path = "tests/data/投标文件template.docx"
    
    if not Path(test_doc_path).exists():
        print(f"❌ 测试文档不存在: {test_doc_path}")
        return
    
    if not Path(template_path).exists():
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    # 创建内容生成器
    generator = ContentGenerator()
    
    try:
        print("⏳ 开始生成投标方案...")
        result = await generator.generate_proposal(
            project=project,
            document_path=test_doc_path,
            template_path=template_path
        )
        
        if result["status"] == "success":
            print("✅ 投标方案生成成功!")
            print(f"📄 文档路径: {result['document_path']}")
            print(f"📋 生成了{len(result['sections'])}个章节")
            
            # 显示部分章节信息
            print("\n📚 章节预览:")
            for i, section in enumerate(result['sections'][:5]):
                print(f"  {i+1}. {section['title']} (Level {section['level']})")
                if section.get('is_leaf'):
                    print(f"     🍃 叶子节点 - 已生成: {section.get('is_generated', False)}")
                else:
                    print(f"     🌿 父节点 - 子节点数: {section.get('children_count', 0)}")
            
            if len(result['sections']) > 5:
                print(f"     ... 还有{len(result['sections']) - 5}个章节")
            
            return result
        else:
            print(f"❌ 生成失败: {result['error']}")
            return None
            
    except Exception as e:
        print(f"❌ 生成异常: {e}")
        import traceback
        traceback.print_exc()
        return None


async def test_template_styles():
    """测试模板样式功能"""
    print("\n🎨 测试模板样式功能...")
    
    from docx import Document
    
    template_path = Path("tests/data/投标文件template.docx")
    if not template_path.exists():
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    # 加载模板
    doc = Document(str(template_path))
    
    print("📋 可用样式:")
    styles = []
    for style in doc.styles:
        if "标书" in style.name:
            styles.append(style.name)
            print(f"  ✅ {style.name}")
    
    # 测试样式应用
    generator = ContentGenerator()
    
    # 测试样式检查方法
    for style_name in ["标书1级", "标书2级", "标书3级", "标书4级", "标书5级", "标书正文"]:
        has_style = generator._has_style(doc, style_name)
        print(f"  {style_name}: {'✅' if has_style else '❌'}")
    
    # 测试级别映射
    for level in range(1, 8):
        style_name = generator._get_title_style(level)
        print(f"  Level {level} -> {style_name}")
    
    return styles


def test_document_creation():
    """测试文档创建功能"""
    print("\n📝 测试文档创建功能...")
    
    from docx import Document
    
    template_path = Path("tests/data/投标文件template.docx")
    if not template_path.exists():
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    # 创建基于模板的文档
    doc = Document(str(template_path))
    
    # 清空现有内容
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
    
    # 添加测试内容
    title_para = doc.add_paragraph("IPTV系统建设项目技术方案")
    title_para.style = "标书1级"
    
    section1_para = doc.add_paragraph("1. 系统架构设计")
    section1_para.style = "标书2级"
    
    subsection_para = doc.add_paragraph("1.1 总体架构")
    subsection_para.style = "标书3级"
    
    content_para = doc.add_paragraph(
        "本系统采用微服务架构设计，具备高可用性和可扩展性。"
        "系统分为前端展示层、业务逻辑层、数据访问层和基础设施层。"
        "各层之间通过标准接口进行通信，确保系统的模块化和可维护性。"
    )
    content_para.style = "标书正文"
    
    # 保存测试文档
    test_output = Path("outputs/test_template_output.docx")
    test_output.parent.mkdir(exist_ok=True)
    doc.save(str(test_output))
    
    print(f"✅ 测试文档已保存: {test_output}")
    return test_output


async def main():
    """主测试函数"""
    print("🤖 AI投标模板样式测试")
    print("=" * 50)
    
    # 测试模板样式
    await test_template_styles()
    
    # 测试文档创建
    test_document_creation()
    
    # 测试完整生成流程
    await test_template_generation()
    
    print("\n" + "=" * 50)
    print("✅ 测试完成!")


if __name__ == "__main__":
    asyncio.run(main())
