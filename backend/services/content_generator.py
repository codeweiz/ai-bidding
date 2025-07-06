import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.models.generation import WorkflowState
from backend.models.project import Project
from backend.services.document_parser import document_parser
from backend.services.workflow_engine import workflow_engine

logger = logging.getLogger(__name__)


class ContentGenerator:
    """内容生成器，负责整个生成流程的协调"""

    def __init__(self):
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
        self.default_template_path = Path("tests/data/投标文件template.docx")

    async def generate_proposal(self, project: Project, document_path: str,
                                template_path: Optional[str] = None) -> Dict[str, Any]:
        """生成完整的投标方案"""
        logger.info(f"开始生成投标方案，项目: {project.name}")

        try:
            # 1. 解析文档
            document_result = document_parser.parse_document(Path(document_path))
            document_content = "\n".join([doc.page_content for doc in document_result["documents"]])

            # 2. 创建工作流状态
            workflow_state = WorkflowState(
                project_id=project.id or "unknown",
                current_step="start",
                document_content=document_content,
                enable_differentiation=project.enable_differentiation
            )

            # 3. 运行工作流
            final_state = await workflow_engine.run_workflow(workflow_state)

            # 处理工作流返回的结果（可能是字典或WorkflowState对象）
            if isinstance(final_state, dict):
                # 如果是字典，转换为WorkflowState对象
                final_state = WorkflowState(**final_state)

            if hasattr(final_state, 'error') and final_state.error:
                return {
                    "status": "error",
                    "error": final_state.error
                }

            # 4. 生成Word文档（使用模板）
            word_doc_path = await self._generate_word_document(project, final_state, template_path)

            return {
                "status": "success",
                "requirements_analysis": final_state.requirements_analysis,
                "outline": final_state.outline,
                "sections": final_state.sections,
                "document_path": str(word_doc_path)
            }

        except Exception as e:
            logger.error(f"生成投标方案失败: {e}")
            return {
                "status": "error",
                "error": str(e)
            }

    async def _generate_word_document(self, project: Project, state: WorkflowState,
                                      template_path: Optional[str] = None) -> Path:
        """生成Word文档 - 使用模板样式"""
        logger.info(f"开始生成Word文档，项目: {project.name}")

        # 确定模板路径
        if template_path and Path(template_path).exists():
            template_doc_path = Path(template_path)
        else:
            template_doc_path = self.default_template_path

        if not template_doc_path.exists():
            logger.warning(f"模板文件不存在: {template_doc_path}，使用空白文档")
            doc = Document()
        else:
            logger.info(f"使用模板: {template_doc_path}")
            doc = Document(str(template_doc_path))

        # 清空模板内容，保留样式
        for paragraph in doc.paragraphs[:]:
            p = paragraph._element
            p.getparent().remove(p)

        # 添加标题
        title_para = doc.add_paragraph(f'{project.name} - 技术方案')
        if self._has_style(doc, "标书1级"):
            title_para.style = "标书1级"
        else:
            title_para.style = "Heading 1"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加目录占位符
        toc_para = doc.add_paragraph('目录')
        if self._has_style(doc, "标书2级"):
            toc_para.style = "标书2级"
        else:
            toc_para.style = "Heading 2"

        doc.add_paragraph('（此处应插入自动生成的目录）')
        doc.add_page_break()

        # 添加章节内容
        for section in state.sections:
            if not section.get("is_generated"):
                continue

            # 添加章节标题 - 使用模板样式
            title_para = doc.add_paragraph(section["title"])
            style_name = self._get_title_style(section["level"])
            if self._has_style(doc, style_name):
                title_para.style = style_name
            else:
                # 回退到标准样式
                title_para.style = f"Heading {min(section['level'], 9)}"

            # 添加章节内容
            content = section.get("differentiated_content") or section.get("content", "")
            if content:
                # 处理内容，支持mermaid代码块
                processed_content = self._process_content_with_diagrams(content)

                # 按段落分割内容
                paragraphs = processed_content.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        # 检查是否是代码块
                        if para.startswith('```') and para.endswith('```'):
                            # 代码块使用等宽字体
                            code_para = doc.add_paragraph(para)
                            code_para.style = "Normal"
                        else:
                            # 正文使用标书正文样式
                            content_para = doc.add_paragraph(para)
                            if self._has_style(doc, "标书正文"):
                                content_para.style = "标书正文"
                            else:
                                content_para.style = "Normal"

        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project.name}_{timestamp}.docx"
        file_path = self.output_dir / filename

        doc.save(str(file_path))
        logger.info(f"Word文档生成完成: {file_path}")

        return file_path

    def _has_style(self, doc: Document, style_name: str) -> bool:
        """检查文档是否有指定样式"""
        try:
            for style in doc.styles:
                if style.name == style_name:
                    return True
            return False
        except:
            return False

    def _get_title_style(self, level: int) -> str:
        """根据级别获取标题样式名称"""
        if level <= 5:
            return f"标书{level}级"
        else:
            return "标书5级"  # 超过5级的都用5级样式

    def _process_content_with_diagrams(self, content: str) -> str:
        """处理内容中的图表代码"""
        # 这里可以后续扩展，将mermaid代码转换为图片
        # 目前先保持原样
        return content

    async def analyze_requirements_only(self, document_path: str) -> Dict[str, Any]:
        """仅分析需求，不生成完整方案"""
        logger.info(f"开始分析需求，文档: {document_path}")

        try:
            # 解析文档
            document_result = document_parser.parse_document(Path(document_path))
            document_content = "\n".join([doc.page_content for doc in document_result["documents"]])

            # 创建简化的工作流状态
            workflow_state = WorkflowState(
                project_id="temp",
                current_step="start",
                document_content=document_content
            )

            # 只运行需求分析步骤
            from backend.services.llm_service import llm_service
            result = await llm_service.analyze_requirements(document_content)

            if result["status"] == "success":
                return {
                    "status": "success",
                    "analysis": result["analysis"],
                    "document_info": {
                        "name": document_result["file_name"],
                        "type": document_result["file_type"],
                        "pages": document_result["metadata"]["total_pages"],
                        "chunks": document_result["metadata"]["total_chunks"]
                    }
                }
            else:
                return {
                    "status": "error",
                    "error": result.get("error", "需求分析失败")
                }

        except Exception as e:
            logger.error(f"需求分析失败: {e}")
            return {
                "status": "error",
                "error": str(e)
            }

    async def generate_outline_only(self, requirements_analysis: str) -> Dict[str, Any]:
        """仅生成提纲"""
        logger.info("开始生成提纲")

        try:
            from backend.services.llm_service import llm_service
            result = await llm_service.generate_outline(requirements_analysis)

            if result["status"] == "success":
                return {
                    "status": "success",
                    "outline": result["outline"]
                }
            else:
                return {
                    "status": "error",
                    "error": result.get("error", "提纲生成失败")
                }

        except Exception as e:
            logger.error(f"提纲生成失败: {e}")
            return {
                "status": "error",
                "error": str(e)
            }

    async def generate_section_content(self, section_title: str, requirements: str,
                                       context: str = "") -> Dict[str, Any]:
        """生成单个章节内容"""
        logger.info(f"开始生成章节内容: {section_title}")

        try:
            from backend.services.llm_service import llm_service
            result = await llm_service.generate_content(section_title, requirements, context)

            if result["status"] == "success":
                return {
                    "status": "success",
                    "content": result["content"]
                }
            else:
                return {
                    "status": "error",
                    "error": result.get("error", "内容生成失败")
                }

        except Exception as e:
            logger.error(f"章节内容生成失败: {e}")
            return {
                "status": "error",
                "error": str(e)
            }

    def get_output_files(self) -> List[Dict[str, Any]]:
        """获取输出文件列表"""
        files = []

        if self.output_dir.exists():
            for file_path in self.output_dir.glob("*.docx"):
                stat = file_path.stat()
                files.append({
                    "name": file_path.name,
                    "path": str(file_path),
                    "size": stat.st_size,
                    "created_at": datetime.fromtimestamp(stat.st_ctime),
                    "modified_at": datetime.fromtimestamp(stat.st_mtime)
                })

        # 按修改时间倒序排列
        files.sort(key=lambda x: x["modified_at"], reverse=True)
        return files


# 全局实例
content_generator = ContentGenerator()
