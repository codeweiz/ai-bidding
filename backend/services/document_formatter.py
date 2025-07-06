"""
文档格式化服务 - 独立的格式化功能
"""
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class DocumentFormatter:
    """文档格式化器 - 独立的格式化服务"""
    
    def __init__(self):
        self.default_template_path = Path("tests/data/投标文件template.docx")
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
        
        # 样式映射配置
        self.style_mapping = {
            1: "标书1级",
            2: "标书2级", 
            3: "标书3级",
            4: "标书4级",
            5: "标书5级",
            "content": "标书正文",
            "title": "标书1级"
        }
        
        # 格式化配置
        self.format_config = {
            "auto_numbering": True,
            "include_toc": True,
            "page_break_after_toc": True,
            "center_title": True,
            "process_diagrams": True
        }
    
    async def format_raw_text(self, raw_text: str, project_name: str = "格式化文档", 
                             template_path: Optional[str] = None) -> Path:
        """格式化原始文本为docx文档"""
        logger.info(f"开始格式化原始文本: {project_name}")
        
        try:
            # 解析原始文本结构
            sections = self._parse_raw_text(raw_text)
            
            # 创建格式化文档
            doc_path = await self._create_formatted_document(
                sections, project_name, template_path
            )
            
            logger.info(f"文本格式化完成: {doc_path}")
            return doc_path
            
        except Exception as e:
            logger.error(f"格式化原始文本失败: {e}")
            raise
    
    async def format_sections_data(self, sections_data: List[Dict[str, Any]], 
                                  project_name: str = "格式化文档",
                                  template_path: Optional[str] = None) -> Path:
        """格式化章节数据为docx文档"""
        logger.info(f"开始格式化章节数据: {project_name}")
        
        try:
            # 创建格式化文档
            doc_path = await self._create_formatted_document(
                sections_data, project_name, template_path
            )
            
            logger.info(f"章节数据格式化完成: {doc_path}")
            return doc_path
            
        except Exception as e:
            logger.error(f"格式化章节数据失败: {e}")
            raise
    
    def _parse_raw_text(self, raw_text: str) -> List[Dict[str, Any]]:
        """解析原始文本为章节结构"""
        sections = []
        lines = raw_text.split('\n')
        current_section = None
        content_buffer = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题行（支持多种格式）
            title_match = self._detect_title(line)
            if title_match:
                # 保存前一个章节
                if current_section and content_buffer:
                    current_section["content"] = "\n".join(content_buffer)
                    sections.append(current_section)
                    content_buffer = []
                
                # 创建新章节
                level, title = title_match
                current_section = {
                    "title": title,
                    "level": level,
                    "content": "",
                    "is_generated": True,
                    "order": len(sections) + 1
                }
            else:
                # 累积内容
                if current_section:
                    content_buffer.append(line)
        
        # 保存最后一个章节
        if current_section and content_buffer:
            current_section["content"] = "\n".join(content_buffer)
            sections.append(current_section)
        
        return sections
    
    def _detect_title(self, line: str) -> Optional[tuple]:
        """检测标题行并返回(level, title)"""
        # 数字编号格式: 1. 1.1 1.1.1 等
        number_pattern = r'^(\d+(?:\.\d+)*)\s+(.+)$'
        match = re.match(number_pattern, line)
        if match:
            number_part = match.group(1)
            title = match.group(2)
            level = len(number_part.split('.'))
            return (level, title)
        
        # Markdown格式: # ## ### 等
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            return (level, title)
        
        # 中文序号格式: 一、二、三、等
        chinese_pattern = r'^[一二三四五六七八九十]+[、．]\s*(.+)$'
        if re.match(chinese_pattern, line):
            title = re.match(chinese_pattern, line).group(1)
            return (1, title)
        
        return None
    
    async def _create_formatted_document(self, sections: List[Dict[str, Any]], 
                                       project_name: str,
                                       template_path: Optional[str] = None) -> Path:
        """创建格式化的Word文档"""
        # 确定模板路径
        if template_path and Path(template_path).exists():
            template_doc_path = Path(template_path)
        else:
            template_doc_path = self.default_template_path
        
        # 创建文档
        if template_doc_path.exists():
            logger.info(f"使用模板: {template_doc_path}")
            doc = Document(str(template_doc_path))
            # 清空模板内容，保留样式
            self._clear_document_content(doc)
        else:
            logger.warning(f"模板文件不存在: {template_doc_path}，使用空白文档")
            doc = Document()
        
        # 添加标题
        if self.format_config["center_title"]:
            title_para = doc.add_paragraph(f'{project_name}')
            if self._has_style(doc, self.style_mapping["title"]):
                title_para.style = self.style_mapping["title"]
            else:
                title_para.style = "Heading 1"
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加目录
        if self.format_config["include_toc"]:
            toc_para = doc.add_paragraph('目录')
            if self._has_style(doc, self.style_mapping[2]):
                toc_para.style = self.style_mapping[2]
            else:
                toc_para.style = "Heading 2"
            
            doc.add_paragraph('（此处应插入自动生成的目录）')
            
            if self.format_config["page_break_after_toc"]:
                doc.add_page_break()
        
        # 添加章节内容
        for section in sections:
            if not section.get("is_generated", True):
                continue
            
            # 添加章节标题
            title_para = doc.add_paragraph(section["title"])
            level = section.get("level", 1)
            style_name = self._get_title_style(level)
            
            if self._has_style(doc, style_name):
                title_para.style = style_name
            else:
                title_para.style = f"Heading {min(level, 9)}"
            
            # 添加章节内容
            content = section.get("content", "")
            if content:
                # 处理内容
                processed_content = self._process_content(content)
                
                # 按段落分割内容
                paragraphs = processed_content.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        self._add_content_paragraph(doc, para)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project_name}_格式化_{timestamp}.docx"
        file_path = self.output_dir / filename
        
        doc.save(str(file_path))
        logger.info(f"格式化文档生成完成: {file_path}")
        
        return file_path
    
    def _clear_document_content(self, doc: Document):
        """清空文档内容，保留样式"""
        # 清空段落
        for paragraph in doc.paragraphs[:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # 清空表格
        for table in doc.tables[:]:
            t = table._element
            t.getparent().remove(t)
    
    def _has_style(self, doc: Document, style_name: str) -> bool:
        """检查文档是否有指定样式"""
        try:
            for style in doc.styles:
                if style.name == style_name:
                    return True
            return False
        except:
            return False
    
    def _get_title_style(self, level: int) -> str:
        """根据级别获取标题样式名称"""
        return self.style_mapping.get(level, self.style_mapping[5])
    
    def _process_content(self, content: str) -> str:
        """处理内容格式"""
        if not self.format_config["process_diagrams"]:
            return content
        
        # 处理图表代码块
        # 这里可以扩展为将mermaid/plantuml转换为图片
        return content
    
    def _add_content_paragraph(self, doc: Document, para_text: str):
        """添加内容段落"""
        # 检查是否是代码块
        if para_text.startswith('```') and para_text.endswith('```'):
            # 代码块使用等宽字体
            code_para = doc.add_paragraph(para_text)
            code_para.style = "Normal"
        else:
            # 正文使用标书正文样式
            content_para = doc.add_paragraph(para_text)
            if self._has_style(doc, self.style_mapping["content"]):
                content_para.style = self.style_mapping["content"]
            else:
                content_para.style = "Normal"
    
    def update_format_config(self, config: Dict[str, Any]):
        """更新格式化配置"""
        self.format_config.update(config)
        logger.info(f"格式化配置已更新: {config}")
    
    def update_style_mapping(self, mapping: Dict[Any, str]):
        """更新样式映射"""
        self.style_mapping.update(mapping)
        logger.info(f"样式映射已更新: {mapping}")


# 全局实例
document_formatter = DocumentFormatter()
