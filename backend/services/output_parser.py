"""
输出解析器 - 负责将AI生成的内容转换为Word格式，确保格式正确
"""
import logging
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """文档章节"""
    title: str
    level: int
    content: str
    subsections: List['DocumentSection'] = None

    def __post_init__(self):
        if self.subsections is None:
            self.subsections = []


@dataclass
class DocumentStyle:
    """文档样式配置"""
    font_name: str = "宋体"
    font_size: int = 12
    line_spacing: float = 1.5
    paragraph_spacing_before: int = 6
    paragraph_spacing_after: int = 6
    heading_font_name: str = "黑体"
    heading_font_size: Dict[int, int] = None

    def __post_init__(self):
        if self.heading_font_size is None:
            self.heading_font_size = {
                1: 16,  # 一级标题
                2: 14,  # 二级标题
                3: 12,  # 三级标题
            }


class OutputParser:
    """输出解析器类"""

    def __init__(self, style: DocumentStyle = None):
        self.style = style or DocumentStyle()
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)

    def parse_markdown_to_sections(self, markdown_content: str) -> List[DocumentSection]:
        """解析Markdown内容为文档章节"""
        logger.info("开始解析Markdown内容")

        sections = []
        current_section = None
        current_content = []

        lines = markdown_content.split('\n')

        for line in lines:
            # 检查是否是标题行
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())

            if heading_match:
                # 保存当前章节
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    sections.append(current_section)

                # 创建新章节
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                current_section = DocumentSection(title=title, level=level, content="")
                current_content = []
            else:
                # 添加到当前章节内容
                if line.strip():  # 忽略空行
                    current_content.append(line)

        # 保存最后一个章节
        if current_section:
            current_section.content = '\n'.join(current_content).strip()
            sections.append(current_section)

        logger.info(f"解析完成，共 {len(sections)} 个章节")
        return sections

    def parse_plain_text_to_sections(self, text_content: str, outline: str = None) -> List[DocumentSection]:
        """解析纯文本内容为文档章节"""
        logger.info("开始解析纯文本内容")

        sections = []

        # 如果有提纲，使用提纲结构
        if outline:
            outline_sections = self._extract_outline_structure(outline)
            content_parts = self._split_content_by_outline(text_content, outline_sections)

            for i, (title, level) in enumerate(outline_sections):
                content = content_parts.get(i, "")
                sections.append(DocumentSection(title=title, level=level, content=content))
        else:
            # 尝试自动识别章节
            sections = self._auto_detect_sections(text_content)

        logger.info(f"解析完成，共 {len(sections)} 个章节")
        return sections

    def _extract_outline_structure(self, outline: str) -> List[Tuple[str, int]]:
        """从提纲中提取结构"""
        structure = []
        lines = outline.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 匹配不同的标题格式
            patterns = [
                (r'^([1-9]\d*)\.\s*(.+)$', 1),  # 1. 标题
                (r'^([1-9]\d*\.[1-9]\d*)\s*(.+)$', 2),  # 1.1 标题
                (r'^([1-9]\d*\.[1-9]\d*\.[1-9]\d*)\s*(.+)$', 3),  # 1.1.1 标题
                (r'^([一二三四五六七八九十]+)、\s*(.+)$', 1),  # 一、标题
                (r'^（([一二三四五六七八九十]+)）\s*(.+)$', 2),  # （一）标题
            ]

            for pattern, level in patterns:
                match = re.match(pattern, line)
                if match:
                    title = match.group(2) if len(match.groups()) > 1 else match.group(1)
                    structure.append((title, level))
                    break

        return structure

    def _split_content_by_outline(self, content: str, outline_sections: List[Tuple[str, int]]) -> Dict[int, str]:
        """根据提纲分割内容"""
        content_parts = {}

        # 简单的内容分割逻辑
        # 实际应用中可能需要更复杂的算法
        paragraphs = content.split('\n\n')
        paragraphs_per_section = max(1, len(paragraphs) // len(outline_sections))

        for i, (title, level) in enumerate(outline_sections):
            start_idx = i * paragraphs_per_section
            end_idx = min((i + 1) * paragraphs_per_section, len(paragraphs))

            section_content = '\n\n'.join(paragraphs[start_idx:end_idx])
            content_parts[i] = section_content

        return content_parts

    def _auto_detect_sections(self, content: str) -> List[DocumentSection]:
        """自动检测章节"""
        sections = []

        # 按段落分割
        paragraphs = content.split('\n\n')

        # 简单的章节检测逻辑
        current_section = DocumentSection(title="主要内容", level=1, content="")
        section_content = []

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # 检查是否可能是标题
            if (len(paragraph) < 50 and
                    not paragraph.endswith('。') and
                    ('方案' in paragraph or '设计' in paragraph or '实现' in paragraph)):

                # 保存当前章节
                if section_content:
                    current_section.content = '\n\n'.join(section_content)
                    sections.append(current_section)

                # 创建新章节
                current_section = DocumentSection(title=paragraph, level=1, content="")
                section_content = []
            else:
                section_content.append(paragraph)

        # 保存最后一个章节
        if section_content:
            current_section.content = '\n\n'.join(section_content)
            sections.append(current_section)

        return sections

    def create_word_document(
            self,
            sections: List[DocumentSection],
            project_name: str,
            metadata: Dict[str, Any] = None
    ) -> Path:
        """创建Word文档"""
        logger.info(f"开始创建Word文档: {project_name}")

        # 创建文档
        doc = Document()

        # 设置文档样式
        self._setup_document_styles(doc)

        # 添加文档标题
        title = doc.add_heading(f'{project_name} - 技术方案', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元数据
        if metadata:
            self._add_metadata_section(doc, metadata)

        # 添加目录占位符
        doc.add_heading('目录', level=1)
        doc.add_paragraph('（此处应插入自动生成的目录）')
        doc.add_page_break()

        # 添加章节内容
        for section in sections:
            self._add_section_to_document(doc, section)

        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project_name}_{timestamp}.docx"
        file_path = self.output_dir / filename

        doc.save(str(file_path))
        logger.info(f"Word文档创建完成: {file_path}")

        return file_path

    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置正文样式
        styles = doc.styles

        # 正文样式
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.style.font_name
        normal_font.size = Pt(self.style.font_size)

        # 段落格式
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing = self.style.line_spacing
        normal_paragraph.space_before = Pt(self.style.paragraph_spacing_before)
        normal_paragraph.space_after = Pt(self.style.paragraph_spacing_after)

        # 标题样式
        for level in range(1, 4):
            heading_style = styles[f'Heading {level}']
            heading_font = heading_style.font
            heading_font.name = self.style.heading_font_name
            heading_font.size = Pt(self.style.heading_font_size.get(level, 12))
            heading_font.bold = True

            # 标题段落格式
            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.space_before = Pt(12)
            heading_paragraph.space_after = Pt(6)

    def _add_metadata_section(self, doc: Document, metadata: Dict[str, Any]):
        """添加元数据章节"""
        doc.add_heading('文档信息', level=1)

        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'

        for key, value in metadata.items():
            row = info_table.add_row()
            row.cells[0].text = str(key)
            row.cells[1].text = str(value)

        doc.add_paragraph()

    def _add_section_to_document(self, doc: Document, section: DocumentSection):
        """添加章节到文档"""
        # 添加标题
        if section.level <= 3:
            doc.add_heading(section.title, level=section.level)
        else:
            # 超过3级的标题作为加粗段落
            para = doc.add_paragraph()
            run = para.add_run(section.title)
            run.bold = True

        # 添加内容
        if section.content:
            # 按段落分割内容
            paragraphs = section.content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    # 处理特殊格式
                    self._add_formatted_paragraph(doc, para_text)

        # 添加子章节
        for subsection in section.subsections:
            self._add_section_to_document(doc, subsection)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """添加格式化段落"""
        # 检查是否是列表项
        if re.match(r'^\s*[•\-\*]\s+', text):
            # 列表项
            para = doc.add_paragraph(text.strip(), style='List Bullet')
        elif re.match(r'^\s*\d+\.\s+', text):
            # 编号列表
            para = doc.add_paragraph(text.strip(), style='List Number')
        else:
            # 普通段落
            para = doc.add_paragraph(text.strip())

        return para


# 全局实例
output_parser = OutputParser()
