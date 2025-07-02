"""
增强的输出解析器 - 支持专业的Word文档格式生成
"""
import re
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """文档章节"""
    title: str
    level: int
    content: str
    section_type: str = "normal"  # "normal", "table", "list", "technical"


@dataclass
class TableData:
    """表格数据"""
    headers: List[str]
    rows: List[List[str]]
    title: Optional[str] = None


class EnhancedOutputParser:
    """增强的输出解析器，支持专业的Word文档格式"""
    
    def __init__(self):
        self.outputs_dir = Path("outputs")
        self.outputs_dir.mkdir(exist_ok=True)

    def create_professional_word_document(
        self,
        sections: List[DocumentSection],
        project_name: str,
        metadata: Dict[str, Any] = None
    ) -> Path:
        """创建专业的Word文档"""
        logger.info(f"开始创建专业Word文档: {project_name}")

        # 创建文档
        doc = Document()
        
        # 设置专业文档样式
        self._setup_professional_styles(doc)
        
        # 添加文档标题页
        self._add_title_page(doc, project_name, metadata)
        
        # 添加目录页
        self._add_table_of_contents(doc, sections)
        
        # 添加章节内容
        for section in sections:
            self._add_professional_section(doc, section)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project_name}_{timestamp}.docx"
        file_path = self.outputs_dir / filename
        
        doc.save(str(file_path))
        logger.info(f"专业Word文档创建完成: {file_path}")
        
        return file_path

    def _setup_professional_styles(self, doc: Document):
        """设置专业文档样式"""
        styles = doc.styles
        
        # 设置标题样式
        for i in range(1, 4):
            heading_style = styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = '微软雅黑'
            heading_font.size = Pt(16 - i * 2)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 0, 0)
            
            # 设置段落格式
            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.space_before = Pt(12)
            heading_paragraph.space_after = Pt(6)
            heading_paragraph.keep_with_next = True
        
        # 设置正文样式
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '宋体'
        normal_font.size = Pt(12)
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing = 1.5
        normal_paragraph.space_after = Pt(6)
        normal_paragraph.first_line_indent = Inches(0.5)

    def _add_title_page(self, doc: Document, project_name: str, metadata: Dict[str, Any]):
        """添加专业标题页"""
        # 主标题
        title = doc.add_heading(project_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # 副标题
        subtitle = doc.add_heading('技术方案书', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.name = '微软雅黑'
        subtitle_run.font.size = Pt(18)
        
        # 添加空行
        for _ in range(5):
            doc.add_paragraph()
        
        # 添加元数据表格
        if metadata:
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            for key, value in metadata.items():
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(value)
                
                # 设置表格样式
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)
        
        # 分页
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, sections: List[DocumentSection]):
        """添加目录页"""
        # 目录标题
        toc_title = doc.add_heading('目录', level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 目录内容
        for section in sections:
            if section.level <= 3:  # 只显示前三级标题
                toc_paragraph = doc.add_paragraph()
                
                # 添加缩进
                indent = (section.level - 1) * 0.5
                toc_paragraph.paragraph_format.left_indent = Inches(indent)
                
                # 添加章节标题
                run = toc_paragraph.add_run(section.title)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                
                # 添加页码占位符（实际项目中可以使用Word的自动目录功能）
                toc_paragraph.add_run('\t...\t1')
        
        # 分页
        doc.add_page_break()

    def _add_professional_section(self, doc: Document, section: DocumentSection):
        """添加专业章节内容"""
        # 添加章节标题
        heading = doc.add_heading(section.title, level=section.level)
        
        # 处理章节内容
        if section.section_type == "table":
            self._add_table_content(doc, section.content)
        elif section.section_type == "list":
            self._add_list_content(doc, section.content)
        elif section.section_type == "technical":
            self._add_technical_content(doc, section.content)
        else:
            self._add_normal_content(doc, section.content)

    def _add_normal_content(self, doc: Document, content: str):
        """添加普通文本内容"""
        # 按段落分割内容
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # 检查是否是列表项
                if re.match(r'^\d+\.', para_text.strip()) or re.match(r'^[•·-]', para_text.strip()):
                    self._add_list_item(doc, para_text.strip())
                else:
                    paragraph = doc.add_paragraph(para_text.strip())
                    self._format_paragraph(paragraph)

    def _add_list_content(self, doc: Document, content: str):
        """添加列表内容"""
        lines = content.split('\n')
        
        for line in lines:
            if line.strip():
                self._add_list_item(doc, line.strip())

    def _add_list_item(self, doc: Document, item_text: str):
        """添加列表项"""
        paragraph = doc.add_paragraph(item_text, style='List Bullet')
        self._format_paragraph(paragraph)

    def _add_technical_content(self, doc: Document, content: str):
        """添加技术内容（可能包含代码块）"""
        # 检查是否包含代码块
        code_pattern = r'```(\w+)?\n(.*?)\n```'
        
        parts = re.split(code_pattern, content, flags=re.DOTALL)
        
        for i, part in enumerate(parts):
            if i % 3 == 0:  # 普通文本
                if part.strip():
                    self._add_normal_content(doc, part)
            elif i % 3 == 2:  # 代码内容
                self._add_code_block(doc, part)

    def _add_code_block(self, doc: Document, code: str):
        """添加代码块"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        
        # 设置代码块样式
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def _add_table_content(self, doc: Document, content: str):
        """添加表格内容"""
        # 简单的表格解析（实际项目中可以更复杂）
        lines = content.split('\n')
        table_data = []
        
        for line in lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    table_data.append(cells)
        
        if table_data:
            # 创建表格
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            
            # 填充表格数据
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_data

    def _format_paragraph(self, paragraph):
        """格式化段落"""
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

    def parse_markdown_to_sections(self, markdown_content: str) -> List[DocumentSection]:
        """解析Markdown内容为文档章节"""
        logger.info("开始解析Markdown内容")

        sections = []
        current_section = None
        current_content = []

        lines = markdown_content.split('\n')

        for line in lines:
            # 检查是否是标题行
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())

            if heading_match:
                # 保存当前章节
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    sections.append(current_section)

                # 创建新章节
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                # 判断章节类型
                section_type = "normal"
                if "架构" in title or "设计" in title or "技术" in title:
                    section_type = "technical"
                elif "表" in title or "清单" in title:
                    section_type = "table"
                elif "要求" in title or "标准" in title:
                    section_type = "list"
                
                current_section = DocumentSection(
                    title=title, 
                    level=level, 
                    content="",
                    section_type=section_type
                )
                current_content = []
            else:
                # 添加到当前章节内容
                if line.strip():
                    current_content.append(line)

        # 保存最后一个章节
        if current_section:
            current_section.content = '\n'.join(current_content).strip()
            sections.append(current_section)

        logger.info(f"解析完成，共 {len(sections)} 个章节")
        return sections

    def parse_plain_text_to_sections(self, text_content: str, outline: str = None) -> List[DocumentSection]:
        """解析纯文本内容为文档章节"""
        logger.info("开始解析纯文本内容")

        sections = []

        # 如果有提纲，使用提纲结构
        if outline:
            outline_sections = self._extract_outline_structure(outline)
            content_parts = self._split_content_by_outline(text_content, outline_sections)

            for i, (title, level) in enumerate(outline_sections):
                content = content_parts.get(i, "")
                
                # 判断章节类型
                section_type = "normal"
                if "架构" in title or "设计" in title or "技术" in title:
                    section_type = "technical"
                elif "表" in title or "清单" in title:
                    section_type = "table"
                elif "要求" in title or "标准" in title:
                    section_type = "list"
                
                sections.append(DocumentSection(
                    title=title, 
                    level=level, 
                    content=content,
                    section_type=section_type
                ))
        else:
            # 尝试自动识别章节
            sections = self._auto_detect_sections(text_content)

        logger.info(f"解析完成，共 {len(sections)} 个章节")
        return sections

    def _extract_outline_structure(self, outline: str) -> List[tuple]:
        """从提纲中提取章节结构"""
        sections = []
        lines = outline.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 匹配不同级别的标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                sections.append((title, level))
            else:
                # 匹配数字编号的标题
                number_match = re.match(r'^(\d+\.)+\s*(.+)$', line)
                if number_match:
                    level = line.count('.') 
                    title = number_match.group(2).strip()
                    sections.append((title, level))
        
        return sections

    def _split_content_by_outline(self, content: str, outline_sections: List[tuple]) -> Dict[int, str]:
        """根据提纲分割内容"""
        content_parts = {}
        
        # 简单的内容分割逻辑
        # 实际项目中可以使用更复杂的算法
        content_lines = content.split('\n')
        lines_per_section = len(content_lines) // max(len(outline_sections), 1)
        
        for i, (title, level) in enumerate(outline_sections):
            start_idx = i * lines_per_section
            end_idx = (i + 1) * lines_per_section if i < len(outline_sections) - 1 else len(content_lines)
            
            section_content = '\n'.join(content_lines[start_idx:end_idx]).strip()
            content_parts[i] = section_content
        
        return content_parts

    def _auto_detect_sections(self, content: str) -> List[DocumentSection]:
        """自动检测章节"""
        sections = []
        
        # 简单的章节检测逻辑
        paragraphs = content.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                title = f"第{i+1}部分"
                sections.append(DocumentSection(
                    title=title,
                    level=1,
                    content=paragraph.strip(),
                    section_type="normal"
                ))
        
        return sections


# 全局实例
enhanced_output_parser = EnhancedOutputParser()
