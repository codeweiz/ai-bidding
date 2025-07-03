def test():
    from docx import Document

    template = Document("./data/投标文件template.docx")
    styles = template.styles

    print("可用的样式:")
    for style in styles:
        print(f"- {style.name}")

    template.add_paragraph("这是一级自定义标题", style="标书1级")
    template.add_paragraph("这是二级自定义标题", style="标书2级")
    template.add_paragraph("这是三级自定义标题", style="标书3级")
    template.add_paragraph("这是四级自定义标题", style="标书4级")
    template.add_paragraph("这是五级自定义标题", style="标书5级")
    template.add_paragraph("这是正文内容，应该会自动使用你在模板里设置的所有格式，比如字体、字号、缩进、行距等。",
                           style="标书正文")

    template.save("./data/test.docx")
    print("测试文档已保存到 ./data/test.docx")

if __name__ == "__main__":
    test()
