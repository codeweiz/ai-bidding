#!/usr/bin/env python3
"""
创建测试用的docx文件
"""

from docx import Document
from pathlib import Path


def create_test_docx():
    """创建测试用的招标文档"""
    print("📝 创建测试招标文档...")
    
    doc = Document()
    
    # 添加标题
    doc.add_heading('智慧城市IPTV系统建设项目招标文件', 0)
    
    # 添加项目概述
    doc.add_heading('一、项目概述', level=1)
    doc.add_paragraph(
        '本项目旨在建设一套完整的智慧城市IPTV管理系统，包括内容管理、用户管理、设备管理等多个子系统。'
        '系统需要具备高可用性、高并发处理能力和良好的扩展性。'
    )
    
    # 添加技术需求
    doc.add_heading('二、技术需求', level=1)
    
    doc.add_heading('1. 系统架构要求', level=2)
    doc.add_paragraph('- 采用微服务架构设计')
    doc.add_paragraph('- 支持云原生部署')
    doc.add_paragraph('- 具备高可用性和可扩展性')
    doc.add_paragraph('- 系统可用性要求达到99.9%以上')
    
    doc.add_heading('2. 功能需求', level=2)
    
    doc.add_heading('2.1 内容管理系统', level=3)
    doc.add_paragraph('- 支持多种视频格式的上传和转码')
    doc.add_paragraph('- 提供内容分类和标签管理')
    doc.add_paragraph('- 支持内容审核和发布流程')
    doc.add_paragraph('- 提供内容统计和分析功能')
    
    doc.add_heading('2.2 用户管理系统', level=3)
    doc.add_paragraph('- 支持用户注册和认证')
    doc.add_paragraph('- 提供用户权限管理')
    doc.add_paragraph('- 支持用户行为分析')
    doc.add_paragraph('- 提供用户服务和支持')
    
    doc.add_heading('2.3 设备管理系统', level=3)
    doc.add_paragraph('- 支持机顶盒设备管理')
    doc.add_paragraph('- 提供设备状态监控')
    doc.add_paragraph('- 支持远程设备控制')
    doc.add_paragraph('- 提供设备故障诊断')
    
    # 添加性能指标
    doc.add_heading('三、性能指标', level=1)
    doc.add_paragraph('- 系统响应时间：≤2秒')
    doc.add_paragraph('- 并发用户数：≥10000')
    doc.add_paragraph('- 视频流处理能力：≥1000路并发')
    doc.add_paragraph('- 存储容量：≥100TB')
    doc.add_paragraph('- 网络带宽：≥10Gbps')
    
    # 添加技术标准
    doc.add_heading('四、技术标准', level=1)
    doc.add_paragraph('- 遵循国家广电总局相关技术标准')
    doc.add_paragraph('- 支持H.264/H.265视频编码')
    doc.add_paragraph('- 符合IPTV行业标准')
    doc.add_paragraph('- 支持IPv6协议')
    doc.add_paragraph('- 遵循信息安全等级保护要求')
    
    # 添加投标要求
    doc.add_heading('五、投标要求', level=1)
    
    doc.add_heading('1. 资质要求', level=2)
    doc.add_paragraph('- 具有软件企业认定证书')
    doc.add_paragraph('- 具有ISO9001质量管理体系认证')
    doc.add_paragraph('- 具有信息安全服务资质证书')
    doc.add_paragraph('- 具有广电行业相关项目经验')
    
    doc.add_heading('2. 技术方案要求', level=2)
    doc.add_paragraph('- 提供详细的系统架构设计')
    doc.add_paragraph('- 提供完整的功能设计方案')
    doc.add_paragraph('- 提供性能优化和安全保障措施')
    doc.add_paragraph('- 提供项目实施计划和风险控制方案')
    
    doc.add_heading('3. 评分标准', level=2)
    doc.add_paragraph('- 技术方案（40分）')
    doc.add_paragraph('- 项目经验（20分）')
    doc.add_paragraph('- 团队实力（20分）')
    doc.add_paragraph('- 价格因素（20分）')
    
    # 保存文档
    output_path = Path("uploads/test_tender_document.docx")
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✅ 测试招标文档已创建: {output_path}")
    return output_path


if __name__ == "__main__":
    create_test_docx()
